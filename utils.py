from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO
import gspread
from oauth2client.service_account import ServiceAccountCredentials
import streamlit as st
from datetime import datetime


def generate_docx(items):
    """
    Generates a Word Document with different layouts based on award type.
    
    For CTO/FSM Coin: 3-column table (Name-Award, Justification, Stats)
    For Other Awards: 2-column table (Name-Award, Justification)
    
    Args:
        items: List of dictionaries with keys: 'rank', 'name', 'text', 'award', 
               and optionally 'ippt', 'bmi', 'atp', 'previous_awards'
    
    Returns:
        bytes: Word document as bytes for download
    """
    doc = Document()
    
    # Set narrow margins for better page utilization
    section = doc.sections[0]
    section.left_margin = Inches(0.5)
    section.right_margin = Inches(0.5)
    section.top_margin = Inches(0.75)
    section.bottom_margin = Inches(0.75)

    # Configure default font style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # Add document title
    title = doc.add_heading('Award Justifications', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()  # Spacing after title

    # Generate table for each entry
    for idx, entry in enumerate(items, 1):
        award_type = entry.get('award', '')
        
        # Determine if this is a CTO/FSM Coin award
        is_cto_fsm = award_type in ["CTO Coin", "FSM Coin"]
        
        if is_cto_fsm and "(CITATION)" not in entry.get('name', ''):
            # === 3-COLUMN LAYOUT FOR CTO/FSM COIN ===
            table = doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            
            # --- COLUMN 1: RANK, NAME & AWARD (20% width) ---
            cell_1 = table.cell(0, 0)
            cell_1.text = f"{entry.get('rank', '')} {entry.get('name', '')} - {award_type}".strip()
            cell_1.width = Inches(1.5)
            
            # Format name-award cell
            for paragraph in cell_1.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(11)
            
            # --- COLUMN 2: JUSTIFICATION TEXT (50% width) ---
            cell_2 = table.cell(0, 1)
            cell_2.text = entry.get('text', '')
            cell_2.width = Inches(3.5)
            
            # Format justification text
            for paragraph in cell_2.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(11)
            
            # --- COLUMN 3: STATS (IPPT, BMI, ATP, AWARDS) (30% width) ---
            cell_3 = table.cell(0, 2)
            cell_3.width = Inches(2.5)
            
            # Build stats text
            stats_lines = []
            
            if entry.get('ippt'):
                stats_lines.append(f"IPPT: {entry.get('ippt')}")
            
            if entry.get('bmi'):
                stats_lines.append(f"BMI: {entry.get('bmi')}")
            
            if entry.get('atp'):
                stats_lines.append(f"ATP: {entry.get('atp')}")
            
            # Add previous awards as bullet points
            if entry.get('previous_awards'):
                stats_lines.append("")  # Empty line before awards
                stats_lines.append("AWARDS:")
                
                # Split by comma and clean up
                awards_list = [award.strip() for award in entry.get('previous_awards').split(',') if award.strip()]
                
                for award in awards_list:
                    stats_lines.append(f"  - {award}")
            
            # Set stats text
            cell_3.text = "\n".join(stats_lines)
            
            # Format stats cell
            for paragraph in cell_3.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(10)
        
        else:
            # === 2-COLUMN LAYOUT FOR OTHER AWARDS ===
            table = doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            
            # --- COLUMN 1: RANK, NAME & AWARD (30% width) ---
            cell_1 = table.cell(0, 0)
            
            # Add award name to display
            if "(CITATION)" in entry.get('name', ''):
                # For citations, keep the original format
                cell_1.text = f"{entry.get('rank', '')} {entry.get('name', '')}".strip()
            else:
                # For regular awards, add award name
                cell_1.text = f"{entry.get('rank', '')} {entry.get('name', '')} - {award_type}".strip()
            
            cell_1.width = Inches(2.0)
            
            # Format name cell (Bold, slightly larger)
            for paragraph in cell_1.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(12)

            # --- COLUMN 2: JUSTIFICATION TEXT (70% width) ---
            cell_2 = table.cell(0, 1)
            cell_2.text = entry.get('text', '')
            cell_2.width = Inches(5.0)
            
            # Format justification text
            for paragraph in cell_2.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.size = Pt(11)
        
        # Add spacing between entries (except after last entry)
        if idx < len(items):
            doc.add_paragraph()

    # Save document to BytesIO buffer
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


def update_sheet(items):
    """
    Updates Google Sheet with award tracking information.
    
    Appends rows with format:
    [RANK, NAME, COY/NODE, AWARD, MONTH OF AWARD, STATUS, PRESENTATION DATE]
    
    Args:
        items: List of dictionaries with keys: 
               'rank', 'name', 'unit', 'award', 'month'
    
    Google Sheet Structure:
    - Column A: RANK
    - Column B: NAME
    - Column C: COY/NODE
    - Column D: AWARD
    - Column E: MONTH OF AWARD
    - Column F: STATUS (auto-filled as "Nominated")
    - Column G: PRESENTATION DATE (left empty for manual entry)
    
    Note:
        Requires 'gcp_service_account' credentials in Streamlit secrets.
        The Google Sheet must be shared with the service account email.
    """
    try:
        # GOOGLE SHEETS CONFIGURATION
        SPREADSHEET_NAME = "NS AWARDS TRACKING"
        WORKSHEET_NAME = "Sheet1" 

        
        # Check if Google Cloud credentials exist
        if "gcp_service_account" not in st.secrets:
            print("INFO: Skipping sheet update - No GCP credentials in secrets.toml")
            return
        
        # Define required scopes for Google Sheets and Drive access
        scope = [
            "https://spreadsheets.google.com/feeds",
            "https://www.googleapis.com/auth/drive"
        ]
        
        # Authenticate using service account from Streamlit secrets
        creds = ServiceAccountCredentials.from_json_keyfile_dict(
            st.secrets["gcp_service_account"],
            scope
        )
        client = gspread.authorize(creds)
        
        # Open the target spreadsheet and worksheet
        spreadsheet = client.open(SPREADSHEET_NAME)
        
        # Option 1: Open by worksheet name
        try:
            sheet = spreadsheet.worksheet(WORKSHEET_NAME)
        except:
            # Option 2: Use first sheet if name not found
            sheet = spreadsheet.sheet1
        
        # Append each entry as a new row
        for item in items:
            # Skip citation entries (they don't need separate tracking)
            if "(CITATION)" in item.get('name', ''):
                continue
            
            # Prepare row data according to tracking format
            row_data = [
                item.get('rank', ''),                    # Column A: RANK
                item.get('name', ''),                    # Column B: NAME
                item.get('unit', ''),                    # Column C: COY/NODE
                item.get('award', ''),                   # Column D: AWARD
                item.get('month', ''),                   # Column E: MONTH OF AWARD
                "NOMINATED",                             # Column F: STATUS (default)
                ""                                       # Column G: PRESENTATION DATE (empty)
            ]
            
            # Append the row to the sheet
            sheet.append_row(row_data)
            
            print(f"✓ Added to tracking: {item.get('rank')} {item.get('name')}")
        
        print(f"SUCCESS: Added {len([i for i in items if '(CITATION)' not in i.get('name', '')])} entries to Google Sheet")
        
    except gspread.exceptions.SpreadsheetNotFound:
        error_msg = f"ERROR: Spreadsheet '{SPREADSHEET_NAME}' not found. Please check the name."
        print(error_msg)
        st.warning(f"⚠️ {error_msg}")
        
    except gspread.exceptions.WorksheetNotFound:
        error_msg = f"ERROR: Worksheet '{WORKSHEET_NAME}' not found in spreadsheet."
        print(error_msg)
        st.warning(f"⚠️ {error_msg}")
        
    except gspread.exceptions.APIError as e:
        error_msg = f"Google Sheets API error: {str(e)}"
        print(f"ERROR: {error_msg}")
        st.warning(f"⚠️ {error_msg}")
        
    except Exception as e:
        print(f"INFO: Sheet update skipped - {str(e)}")
        # Don't show warning for general errors to avoid disrupting user flow